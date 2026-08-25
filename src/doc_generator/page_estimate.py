"""
Rough DOCX page-count estimator.

Neither LibreOffice nor a scriptable PDF-export path is reliably
available in this environment, so an authoritative page count (as a
real layout engine like Word would report) can't be computed here.
This produces a heuristic estimate instead, from the document's own
page geometry (via python-docx) and the character volume of its
paragraphs/tables — good enough to catch "this report grew back to
5 pages" during development. Treat the result as approximate; verify
the exact count by opening the generated file in Word when it
matters.
"""

from __future__ import annotations

import math
from pathlib import Path

import docx

# Rough average glyph width for a Calibri-like sans body font, as a
# fraction of the font's point size — a standard typographic rule of
# thumb, not a font-metrics lookup.
_CHARS_PER_POINT = 0.52

# Line height (including normal paragraph spacing) as a multiple of
# the font's point size.
_LINE_SPACING_FACTOR = 1.35


def estimate_pages(path: str | Path, *, font_size_pt: float = 11.0) -> float:
    """
    Estimated page count for the DOCX at `path`, given the document's
    own page size/margins and a heuristic text-wrap model. Approximate
    — not a substitute for opening the file in a real word processor.
    """

    document = docx.Document(str(path))
    section = document.sections[0]

    usable_width_pt = section.page_width.pt - section.left_margin.pt - section.right_margin.pt
    usable_height_pt = section.page_height.pt - section.top_margin.pt - section.bottom_margin.pt

    chars_per_line = max(1, int(usable_width_pt / (font_size_pt * _CHARS_PER_POINT)))
    line_height_pt = font_size_pt * _LINE_SPACING_FACTOR
    lines_per_page = max(1.0, usable_height_pt / line_height_pt)

    total_lines = 0.0

    for paragraph in document.paragraphs:

        text = paragraph.text.strip()

        if not text:
            total_lines += 0.4
            continue

        is_heading = paragraph.style is not None and "Heading" in (paragraph.style.name or "")
        is_title = paragraph.style is not None and "Title" in (paragraph.style.name or "")

        wrapped_lines = math.ceil(len(text) / chars_per_line)

        # Headings/titles carry extra space-before/after in the
        # template's styling, so they cost more vertical room than
        # their wrapped-line count alone suggests.
        weight = 2.2 if is_title else (1.7 if is_heading else 1.0)

        total_lines += wrapped_lines * weight

    for table in document.tables:

        for row in table.rows:

            cell_count = max(1, len(row.cells))
            chars_per_cell_line = max(1, chars_per_line // cell_count)

            cell_line_counts = [
                math.ceil(len(cell.text) / chars_per_cell_line) if cell.text.strip() else 1
                for cell in row.cells
            ]

            # Cells render side-by-side; a row's height is set by its
            # tallest cell.
            total_lines += max(cell_line_counts, default=1)

    return round(total_lines / lines_per_page, 2)


def format_report(path: str | Path, *, page_limit: float = 2.0) -> str:
    """A short human-readable estimate line, e.g. for a test/dev script."""

    pages = estimate_pages(path)
    status = "within limit" if pages <= page_limit else "OVER LIMIT"

    return f"Estimated page count: {pages} (limit {page_limit}) — {status}"
