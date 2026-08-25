"""
SAE Case Summary document builder.

Renders the pipeline's structured `CaseSummary` and AI-generated
`ClinicalReport` as a "Clinical Case Narrative" — prose paragraphs
organized chronologically under fixed headings, not field/value
tables.

Narrative text is synthesized here, in plain Python, directly from
already-extracted structured data (`CaseSummary`) and the LLM's own
already-generated narrative field (`ClinicalReport.executive_summary`
/ `.final_outcome`, used as optional lead-ins). No LLM call is made
from this module and no extraction/reasoning logic is touched — this
is presentation only. A clause is written only when the underlying
field is actually present; nothing here invents a clinical fact, a
date, a diagnosis, or an outcome, and a document with fewer records
(or none at all) produces a correspondingly shorter narrative rather
than a differently-shaped one.
"""

from __future__ import annotations

import re
from datetime import date, datetime
from pathlib import Path
from typing import Any

from docx.document import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from models.case_summary import CaseSummary
from models.dcm import DCMData
from models.lab import LabData
from models.maternal import MaternalData
from models.nss import NSSData

from .template import create_document

try:
    from reasoning.report import ClinicalReport
except ImportError:  # pragma: no cover - reasoning package optional at import time
    ClinicalReport = None  # type: ignore[assignment, misc]


# --------------------------------------------------------------
# Value formatting
# --------------------------------------------------------------

# Known source-form sentinel values that must never be presented as a
# genuine clinical measurement. NSSData.respiratory_rate's own
# docstring documents 9999 as meaning "ventilated / not assessable",
# not an actual breaths-per-minute reading.
_SENTINELS: dict[str, set[float]] = {
    "respiratory_rate": {9999},
}

# REDCap records single/multi-select answers bilingually, e.g.
# "जीवित, Alive", "साफ (Clear)", "सामान्य / कोई जटिलता नहीं (Uneventful)".
# Both forms carry the exact same clinical meaning (the English form
# is the study's own translation of the Devanagari one, not a
# separate fact) — showing both in an English-language report reads
# as broken text throughout every section. The English form is
# always already present verbatim in the extracted string; this only
# selects it, it never translates or invents anything.
_DEVANAGARI_RE = re.compile(r"[ऀ-ॿ]")
_TRAILING_PAREN_RE = re.compile(r"\(([^()]*)\)\s*$")


def _english_only(text: str) -> str:

    if not _DEVANAGARI_RE.search(text):
        return text

    paren = _TRAILING_PAREN_RE.search(text)

    if paren and not _DEVANAGARI_RE.search(paren.group(1)):
        candidate = paren.group(1).strip()
        if candidate:
            return candidate

    if "," in text:
        tail = text.rsplit(",", 1)[-1].strip()
        if tail and not _DEVANAGARI_RE.search(tail):
            return tail

    # No isolatable English portion (e.g. a Hindi-only remark with no
    # translation) — strip the Devanagari run and leftover connector
    # punctuation rather than show mixed script, but never return an
    # empty string; fall back to the untouched original in that case.
    stripped = _DEVANAGARI_RE.sub("", text)
    stripped = re.sub(r"[,/()]+", " ", stripped)
    stripped = re.sub(r"\s+", " ", stripped).strip()

    return stripped or text


def _clean(value: Any) -> str | None:
    """A present, non-empty display string for `value`, or None."""

    if value is None:
        return None

    if isinstance(value, bool):
        return "yes" if value else "no"

    if isinstance(value, str):
        stripped = value.strip()
        return _english_only(stripped) if stripped else None

    if isinstance(value, float):
        # A whole-number reading (e.g. weight 1490.0) reads more
        # naturally in prose as "1490" than "1490.0"; a genuine
        # decimal reading (e.g. temperature 37.2) is left as-is.
        return str(int(value)) if value.is_integer() else str(value)

    return str(value)


def _field(record: Any, field_name: str) -> str | None:
    """
    `_clean(record.<field_name>)`, but returns None for a value that
    is a documented sentinel/placeholder rather than a real reading.
    """

    if record is None:
        return None

    value = getattr(record, field_name, None)

    if value is not None and value in _SENTINELS.get(field_name, set()):
        return None

    return _clean(value)


_DATE_FORMATS = (
    "%d-%m-%Y %H:%M",
    "%d-%m-%Y",
    "%d/%m/%Y %H:%M",
    "%d/%m/%Y",
    "%d-%b-%Y %H:%M",
    "%d-%b-%Y",
)


def _parse_date(value: str | None) -> date | None:

    if not value:
        return None

    text = value.strip()

    for fmt in _DATE_FORMATS:
        try:
            return datetime.strptime(text, fmt).date()
        except ValueError:
            continue

    return None


# --------------------------------------------------------------
# OCR / extraction artifact cleanup for free-text remarks
# --------------------------------------------------------------

_SENTENCE_END = (".", "!", "?", "।")


def _dedupe_repeated_phrase(text: str) -> str:
    """
    Collapses a phrase that repeats immediately back-to-back at the
    start of the text — an observed extraction artifact where a DCM
    and NSS remark are near-duplicates and end up concatenated, e.g.
    "Baby severe Baby severe today" -> "Baby severe today".
    """

    words = text.split()
    n = len(words)

    for k in range(n // 2, 0, -1):
        if words[:k] == words[k:2 * k]:
            return " ".join(words[:k] + words[2 * k:])

    return text


# A REDCap question header leaking into a free-text field looks like
# "22- Please enter..." / "20-Is the..." / "25-In your opinion...":
# digit(s), then a dash or comma, then a capitalized word. Two or
# more of these inside one field is a reliable sign that raw,
# scrambled form text was captured instead of the intended answer
# (a known artifact on badly-OCR'd multi-column source pages) — the
# whole field is unsafe to reproduce, not just a trailing fragment.
_FORM_QUESTION_FRAGMENT = re.compile(r"\b\d{1,2}(?:\.\d+)?\s*[-,]\s*[A-Z]")

# Two further, independent signs that a field-boundary match failed
# upstream and a free-text field (diagnosis, event description, cause
# of death, remarks) swallowed unrelated content past its intended
# end: (1) severely fragmented OCR reads as a run of isolated
# single-letter tokens ("t it d i t di", "i j d l 0 2 l /h b i f") that
# never appears in genuine prose beyond the pronouns "a"/"I"; (2) a
# medication-administration list ("inj ... inj ... inj ...") bleeding
# into a field that isn't the medications field.
_ISOLATED_LETTER_EXCEPTIONS = {"a", "i"}


def _isolated_letter_count(text: str) -> int:

    count = 0

    for word in text.split():

        core = word.strip(".,;:")

        if len(core) == 1 and core.isalpha() and core.lower() not in _ISOLATED_LETTER_EXCEPTIONS:
            count += 1

    return count


_MEDICATION_TOKEN = re.compile(r"\binj\b", flags=re.IGNORECASE)


def _is_garbled_form_dump(text: str) -> bool:

    if len(_FORM_QUESTION_FRAGMENT.findall(text)) >= 2:
        return True

    if _isolated_letter_count(text) >= 3:
        return True

    if len(_MEDICATION_TOKEN.findall(text)) >= 3:
        return True

    return False


def _clean_remark(text: str | None) -> str | None:
    """
    Cleanup for a free-text remark/description before it is embedded
    in the narrative: omits the value entirely if it looks like a
    scrambled dump of raw form text; otherwise collapses an
    immediately repeated phrase, and trims a trailing one-or-two
    -character fragment with no sentence-ending punctuation — the
    signature of a mid-word OCR truncation (e.g. "...severe today L").
    This never invents a completion; a cleaned remark is only ever
    shorter than the source, never rewritten or extended.
    """

    cleaned = _clean(text)

    if not cleaned:
        return None

    if _is_garbled_form_dump(cleaned):
        return None

    cleaned = _dedupe_repeated_phrase(cleaned)

    words = cleaned.split()

    if words:
        last = words[-1]
        core = last.rstrip(".,;:")
        if len(core) <= 2 and last[-1:] not in _SENTENCE_END:
            words = words[:-1]

    cleaned = " ".join(words).strip()

    return cleaned if len(cleaned) >= 3 else None


# --------------------------------------------------------------
# Document primitives
# --------------------------------------------------------------


def _join_sentences(sentences: list[str]) -> str:
    """
    Joins clause/sentence fragments into one paragraph, ensuring each
    ends with terminal punctuation first — a fragment sourced
    verbatim from extracted data (e.g. an event description) may not
    end in a period, and joining it directly onto the next sentence
    would otherwise run the two together.
    """

    finished = []

    for sentence in sentences:

        text = sentence.strip()

        if not text:
            continue

        if text[-1] not in _SENTENCE_END:
            text += "."

        finished.append(text)

    return " ".join(finished)


def _paragraph(document: DocxDocument, text: str, *, italic: bool = False):

    paragraph = document.add_paragraph(text)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.08

    if italic and paragraph.runs:
        paragraph.runs[0].italic = True

    return paragraph


def _section_heading(document: DocxDocument, text: str) -> None:

    document.add_heading(text, level=1)


# --------------------------------------------------------------
# Title / identification
# --------------------------------------------------------------


def _title_block(document: DocxDocument, summary: CaseSummary) -> None:

    sae = summary.sae
    maternal = summary.maternal

    hospital = _field(sae, "hospital_name") or _field(maternal, "hospital_name")
    record_id = _field(sae, "record_id") or _clean(summary.case_id)

    title_parts = ["ICMR Emollient Study"]

    if hospital:
        title_parts.append(hospital)

    if record_id:
        title_parts.append(f"Record ID: {record_id}")

    title = document.add_heading(" | ".join(title_parts), level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = document.add_paragraph("Clinical Case Narrative")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(8)

    if subtitle.runs:
        subtitle.runs[0].bold = True
        subtitle.runs[0].font.size = Pt(13)


def _identification_block(document: DocxDocument, summary: CaseSummary) -> None:

    sae = summary.sae
    maternal = summary.maternal

    baby_ref = _field(sae, "patient_name")

    if not baby_ref and maternal and _field(maternal, "mother_name"):
        baby_ref = f"Baby of {_field(maternal, 'mother_name')}"

    hospital = _field(sae, "hospital_name") or _field(maternal, "hospital_name")

    line1_parts = [part for part in [baby_ref, "ICMR Emollient Study", hospital] if part]

    if line1_parts:
        _paragraph(document, "  |  ".join(line1_parts), italic=True)

    record_id = _field(sae, "record_id") or _clean(summary.case_id)
    uid = _field(maternal, "baby_uid")

    line2_parts = []

    if record_id:
        line2_parts.append(f"Record ID: {record_id}")

    if uid:
        line2_parts.append(f"UID: {uid}")

    if line2_parts:
        _paragraph(document, "  |  ".join(line2_parts), italic=True)


# --------------------------------------------------------------
# Section 1 — Case Summary
# --------------------------------------------------------------


def _section_case_summary(document: DocxDocument, summary: CaseSummary) -> None:
    """
    One concise paragraph of key patient/admission facts — the parts
    not already shown in the compact identification lines above
    (name/hospital/record ID/UID): sex, age, date of birth when
    documented, and admission/discharge dates. This is the one
    primary home for admission logistics, so the Outcome section
    doesn't also need to restate them.
    """

    _section_heading(document, "1. Case Summary")

    sae = summary.sae
    maternal = summary.maternal

    baby_ref = _field(sae, "patient_name")

    if not baby_ref and maternal and _field(maternal, "mother_name"):
        baby_ref = f"Baby of {_field(maternal, 'mother_name')}"

    hospital = _field(sae, "hospital_name") or _field(maternal, "hospital_name")

    bits = []

    if _field(sae, "sex"):
        bits.append(f"sex {_field(sae, 'sex')}")

    if _field(sae, "date_of_birth"):
        bits.append(f"date of birth {_field(sae, 'date_of_birth')}")

    if sae and sae.age_days is not None:
        bits.append(f"age {sae.age_days} days")

    opening = baby_ref or "This case"

    if bits:
        opening += " (" + _list_join(bits) + ")"

    if hospital:
        opening += f" was managed at {hospital}"

    sentences = [opening.strip() + "."]

    admin_bits = []

    if _field(sae, "admission_date"):
        admin_bits.append(f"admitted on {_field(sae, 'admission_date')}")

    if _field(sae, "discharge_date"):
        admin_bits.append(f"discharged on {_field(sae, 'discharge_date')}")

    if admin_bits:
        sentences.append("The infant was " + _list_join(admin_bits) + ".")

    _paragraph(document, _join_sentences(sentences))


# --------------------------------------------------------------
# Section 2 — Maternal History
# --------------------------------------------------------------


def _maternal_narrative(maternal: MaternalData) -> str | None:
    """
    Synthesizes maternal history into flowing prose from whichever
    fields are actually present. Built deterministically (rather than
    relying solely on the LLM's narrative) so that every available
    field is guaranteed to be represented, for any future document.
    """

    sentences: list[str] = []

    name = _field(maternal, "mother_name")

    parity_term = None
    if maternal.gravida == 1:
        parity_term = "primigravida"
    elif maternal.gravida is not None and maternal.gravida > 1:
        parity_term = "multigravida"

    gp_bits = []
    if maternal.gravida is not None:
        gp_bits.append(f"G{maternal.gravida}")
    if maternal.para is not None:
        gp_bits.append(f"P{maternal.para}")
    gp_text = "".join(gp_bits) or None

    if name:
        opening = name
        if maternal.mother_age is not None:
            opening += f", a {maternal.mother_age}-year-old"
            if parity_term:
                opening += f" {parity_term}"
            opening += " mother"
        elif parity_term:
            opening += f", a {parity_term} mother"
    elif parity_term:
        opening = f"The mother, a {parity_term},"
    else:
        opening = "The mother"

    if gp_text:
        opening += f" ({gp_text})"

    father = _field(maternal, "father_name")
    if father:
        opening += f", father {father},"

    hospital = _field(maternal, "hospital_name")
    if hospital:
        opening += f" was managed at {hospital}"

    sentences.append(opening.strip().rstrip(",") + ".")

    if maternal.corticosteroids_given is True:
        steroid = "Antenatal corticosteroids were administered"
        details = []
        if _field(maternal, "steroid_type"):
            details.append(_field(maternal, "steroid_type"))
        if maternal.steroid_doses is not None:
            details.append(f"{maternal.steroid_doses} doses")
        if details:
            steroid += " (" + ", ".join(details) + ")"
        sentences.append(steroid + ".")
    elif maternal.corticosteroids_given is False:
        sentences.append("Antenatal corticosteroids were not administered.")

    if _field(maternal, "fever_history"):
        sentences.append(
            f"Fever history during pregnancy/labour was documented as "
            f"{_field(maternal, 'fever_history')}."
        )

    risk_bits = []

    if maternal.foul_smelling_discharge is True:
        risk_bits.append("foul-smelling vaginal discharge")

    if maternal.uterine_tenderness is True:
        risk_bits.append("uterine tenderness")

    if maternal.pprom_over_24h is True:
        risk_bits.append("leaking per vaginum for more than 24 hours (PPROM >24h)")

    if maternal.foetal_distress is True:
        risk_bits.append("foetal distress")

    if risk_bits:
        sentences.append("Antenatal/intrapartum findings included " + ", ".join(risk_bits) + ".")

    if _field(maternal, "amniotic_fluid_appearance"):
        sentences.append(
            f"Amniotic fluid at delivery was {_field(maternal, 'amniotic_fluid_appearance')}."
        )

    labour_bits = []

    if _field(maternal, "labour_type"):
        labour_bits.append(f"labour type {_field(maternal, 'labour_type')}")

    if _field(maternal, "labour_course"):
        labour_bits.append(f"course {_field(maternal, 'labour_course')}")

    if labour_bits:
        sentences.append("Labour: " + ", ".join(labour_bits) + ".")

    delivery_bits = []

    if _field(maternal, "delivery_type"):
        delivery_bits.append(f"delivered by {_field(maternal, 'delivery_type')}")
    elif _field(maternal, "delivery_mode"):
        delivery_bits.append(f"delivered by {_field(maternal, 'delivery_mode')}")

    if _field(maternal, "presentation"):
        delivery_bits.append(f"{_field(maternal, 'presentation')} presentation")

    if _field(maternal, "delivery_attendant"):
        delivery_bits.append(f"assisted by {_field(maternal, 'delivery_attendant')}")

    if delivery_bits:
        sentences.append("The baby was " + ", ".join(delivery_bits) + ".")

    if _field(maternal, "gestational_age"):
        sentences.append(f"Gestational age was documented as {_field(maternal, 'gestational_age')}.")

    screening_bits = []

    if _field(maternal, "blood_group"):
        screening_bits.append(f"blood group {_field(maternal, 'blood_group')}")

    if _field(maternal, "vdrl"):
        screening_bits.append(f"VDRL {_field(maternal, 'vdrl')}")

    if _field(maternal, "hbsag"):
        screening_bits.append(f"HBsAg {_field(maternal, 'hbsag')}")

    if _field(maternal, "hiv_status"):
        screening_bits.append(f"HIV status {_field(maternal, 'hiv_status')}")

    if maternal.pih is True:
        screening_bits.append("pregnancy-induced hypertension")

    if maternal.gdm is True:
        screening_bits.append("gestational diabetes mellitus")

    if _field(maternal, "thyroid_status"):
        screening_bits.append(f"thyroid status {_field(maternal, 'thyroid_status')}")

    if maternal.oligohydramnios is True:
        screening_bits.append("oligohydramnios")

    if _field(maternal, "amniotic_fluid_volume"):
        screening_bits.append(f"amniotic fluid volume {_field(maternal, 'amniotic_fluid_volume')}")

    if screening_bits:
        sentences.append("Antenatal screening/status: " + ", ".join(screening_bits) + ".")

    other_bits = []

    if _field(maternal, "illness_history"):
        other_bits.append(f"illness history: {_field(maternal, 'illness_history')}")

    if _field(maternal, "epilepsy_radiation_drug_history"):
        other_bits.append(
            f"epilepsy/radiation/drug history: {_field(maternal, 'epilepsy_radiation_drug_history')}"
        )

    if maternal.aph is True:
        other_bits.append("antepartum haemorrhage")

    if other_bits:
        sentences.append("Other relevant history: " + "; ".join(other_bits) + ".")

    remark = _clean_remark(_field(maternal, "maternal_remarks"))

    if maternal.significant_maternal_info_flag is True:
        sentence = "Significant maternal information was flagged"
        sentence += f": {remark}." if remark else "."
        sentences.append(sentence)
    elif remark:
        sentences.append(remark)

    return _join_sentences(sentences) if sentences else None


def _section_maternal_history(document: DocxDocument, summary: CaseSummary, clinical_report) -> None:

    _section_heading(document, "2. Maternal History")

    narrative = _maternal_narrative(summary.maternal) if summary.maternal else None

    if not narrative and clinical_report:
        narrative = _clean(getattr(clinical_report, "maternal_history", None))

    _paragraph(document, narrative or "No maternal history data was available for this case.")


# --------------------------------------------------------------
# Section 2 — Day-wise Clinical Course
# --------------------------------------------------------------


def _group_by_day(
    dcm_records: list[DCMData],
    nss_records: list[NSSData],
) -> list[tuple[int, str | None, DCMData | None, NSSData | None]]:
    """
    Groups DCM and NSS records that fall on the same calendar day,
    sorted chronologically, and numbered Day 1, Day 2, ... Records
    whose date cannot be parsed are still included, ordered after
    every dated record, using their own raw date string as the
    grouping key so they are never silently dropped. The number of
    days produced is entirely a function of the input records.
    """

    days: dict[str, dict[str, Any]] = {}

    def register(records: list, date_field: str, kind: str) -> None:

        for record in records:

            raw_date = getattr(record, date_field, None)
            parsed = _parse_date(raw_date)
            key = parsed.isoformat() if parsed else f"unparsed::{raw_date or id(record)}"

            bucket = days.setdefault(
                key, {"dcm": None, "nss": None, "display_date": raw_date, "sort_date": parsed}
            )
            bucket[kind] = record

            if raw_date and not bucket.get("display_date"):
                bucket["display_date"] = raw_date

    register(dcm_records, "monitoring_date", "dcm")
    register(nss_records, "assessment_date", "nss")

    def sort_key(item: tuple[str, dict[str, Any]]):
        _, bucket = item
        sort_date = bucket["sort_date"]
        return (0, sort_date) if sort_date else (1, bucket["display_date"] or "")

    ordered = sorted(days.items(), key=sort_key)

    return [
        (index, bucket["display_date"], bucket["dcm"], bucket["nss"])
        for index, (_, bucket) in enumerate(ordered, start=1)
    ]


def _ventilation_level(nss: NSSData | None) -> int | None:

    if nss is None:
        return None

    if nss.on_invasive_ventilator:
        return 2

    if nss.on_noninvasive_ventilator:
        return 1

    if nss.on_invasive_ventilator is False and nss.on_noninvasive_ventilator is False:
        return 0

    return None


_SEPSIS_PATTERNS: list[tuple[int, tuple[str, ...]]] = [
    (2, ("confirmed sepsis",)),
    (1, ("probable sepsis", "suspected sepsis")),
    (0, ("no signs of sepsis", "no sign of sepsis", "no evidence of sepsis")),
]

_SEPSIS_LABELS = {0: "no signs of sepsis", 1: "probable sepsis", 2: "confirmed sepsis"}


def _sepsis_rank(text: str | None) -> int | None:

    if not text:
        return None

    lowered = text.lower()

    for rank, keywords in _SEPSIS_PATTERNS:
        if any(keyword in lowered for keyword in keywords):
            return rank

    return None


# Compact labels for table cells (the narrative sentence forms above
# read naturally in prose but are too verbose repeated in every row).
_VENTILATION_LABELS_SHORT = {
    0: "Not on ventilatory support",
    1: "On non-invasive ventilatory support",
    2: "On invasive mechanical ventilation",
}


# --------------------------------------------------------------
# Section 3 — Clinical Course (chronology + progression, one narrative)
# --------------------------------------------------------------


# The same abnormal-vitals flags already shown per-day in the NSS
# table cell (see _nss_cell_lines) — reused here to detect when one
# becomes newly present, rather than restating it on every day it
# persists.
_ABNORMAL_VITAL_FLAGS: list[tuple[str, str]] = [
    ("tachycardia", "tachycardia"),
    ("bradycardia", "bradycardia"),
    ("fast_breathing", "fast breathing"),
    ("bradypnea", "bradypnea"),
    ("high_temperature_flag", "high temperature"),
    ("low_temperature_flag", "low body temperature"),
    ("bleeding_observed", "bleeding"),
    ("fontanelle_bulging", "bulging fontanelle"),
]

# Strips dose/route/frequency tokens so the same drug recorded with
# slightly different formatting on two different days ("Inj Capnea
# 08 Mg OD" vs "inj-Capnea 8mg OD") compares as unchanged, while a
# genuinely different drug does not. Best-effort text normalization
# only — used solely to decide whether the *regimen* changed; the
# medication text shown in the day-wise table and in any reported
# event is always the original, unaltered source text.
_MEDICATION_NORMALIZE_RE = re.compile(
    r"[\d.]+\s*(?:mg|ml|mcg|units?|iu)\b"
    r"|\b(?:od|bd|tds|qid|iv|im|po|stat|hrly|hr)\b"
    # Route/formulation prefixes recorded inconsistently between days
    # for the same drug (e.g. "Capnea 8mg" one day, "Inj Capnea 08 Mg
    # OD" the next) — stripped so the drug name itself is what's
    # compared, not whether a route was noted that day.
    r"|\b(?:inj|tab|syp|susp|cap|drops)\b",
    flags=re.IGNORECASE,
)


def _medication_items(text: str | None) -> dict[str, str]:
    """
    Maps a normalized comparison key -> the original item text, for
    each medication entry in a free-text medications field.
    """

    if not text:
        return {}

    items: dict[str, str] = {}

    for raw in re.split(r"[,;]+", text):

        raw = raw.strip(" ,;")

        if not raw:
            continue

        stem = _MEDICATION_NORMALIZE_RE.sub("", raw)
        stem = re.sub(r"[^A-Za-z ]", " ", stem).lower()
        stem = re.sub(r"\s+", " ", stem).strip()

        if stem:
            items[stem] = raw

    return items


def _list_join(items: list[str], conjunction: str = "and") -> str:
    """
    Natural-English list join: "A" / "A and B" / "A, B, and C" —
    instead of a bare comma- or semicolon-separated dump with no
    closing conjunction.
    """

    items = list(items)

    if not items:
        return ""

    if len(items) == 1:
        return items[0]

    if len(items) == 2:
        return f"{items[0]} {conjunction} {items[1]}"

    return ", ".join(items[:-1]) + f", {conjunction} {items[-1]}"


def _clause_join(clauses: list[str]) -> str:
    """
    Joins several independent clinical clauses documented on the same
    day into one sentence — semicolon-separated (standard usage for a
    list whose items are themselves multi-word clauses), with "and"
    before the final clause rather than leaving the list open-ended.
    """

    clauses = list(clauses)

    if not clauses:
        return ""

    if len(clauses) <= 2:
        return "; ".join(clauses)

    return "; ".join(clauses[:-1]) + f"; and {clauses[-1]}"


def _clinical_course_narrative(days: list) -> str | None:
    """
    Synthesizes the hospital/NICU course into a small number of
    connected paragraph sentences — deterministically, from fields
    already present on a DCMData/NSSData record. This is the single
    home for chronology and meaningful clinical change: it does not
    restate every day, only a day on which something actually
    changed (respiratory support escalation, new-onset abnormal
    vitals, a change in documented general condition, sepsis
    -impression progression, a treatment/medication regimen change, a
    clinically noticeable weight change, or death).

    A day whose *only* change is a further sepsis-impression
    progression is not given its own "On Day N" sentence — chaining
    a new dateline for every such day is exactly the fragment-chain
    style this format replaces. Instead it is folded onto the
    previous sentence with "while ...", so multi-step sepsis
    progressions (e.g. no signs -> probable -> confirmed) still each
    appear, in order, without a wall of near-identical short
    sentences. Nothing here infers a diagnosis, a cause, or a
    relationship not already stated in the source records.
    """

    if not days:
        return None

    sentences: list[str] = []

    _, first_date, first_dcm, first_nss = days[0]

    opening_bits = []

    if first_dcm and _field(first_dcm, "vital_status"):
        opening_bits.append(_field(first_dcm, "vital_status").lower())

    if first_dcm and first_dcm.weight is not None:
        opening_bits.append(f"weight {_clean(first_dcm.weight)} g")

    if first_nss and _field(first_nss, "general_condition"):
        opening_bits.append(f"general condition {_field(first_nss, 'general_condition').lower()}")

    first_sepsis_text = _field(first_nss, "sepsis_impression") if first_nss else None

    if first_sepsis_text:
        opening_bits.append(f"sepsis impression {first_sepsis_text.lower()}")

    if opening_bits:
        sentences.append("The infant was initially documented as " + _list_join(opening_bits) + ".")

    previous_ventilation = _ventilation_level(first_nss)
    previous_sepsis_rank = _sepsis_rank(first_sepsis_text)
    previous_weight = first_dcm.weight if first_dcm else None
    previous_general_condition = _field(first_nss, "general_condition") if first_nss else None
    previous_abnormal = {
        label for flag, label in _ABNORMAL_VITAL_FLAGS if first_nss and getattr(first_nss, flag, None)
    }
    previous_medications = _medication_items(_field(first_dcm, "medications")) if first_dcm else {}

    last_day_number = days[-1][0]

    # A sepsis-only day is folded onto the *previous per-day event
    # sentence*, never onto the opening "initially documented"
    # sentence (which describes a starting state, not a transition,
    # and has no date of its own to anchor a "while ..." clause to).
    # `event_sentence_started` tracks whether such a sentence exists
    # yet; `last_append_was_merge` varies the connector so a run of
    # several consecutive sepsis-only days doesn't read as
    # "while ..., while ..., while ...".
    event_sentence_started = False
    last_append_was_merge = False

    for day_number, display_date, dcm, nss in days[1:]:

        is_last_day = day_number == last_day_number

        clauses: list[str] = []

        # Respiratory support escalation.
        ventilation_level = _ventilation_level(nss)

        if (
            ventilation_level is not None
            and previous_ventilation is not None
            and ventilation_level > previous_ventilation
        ):
            support_label = _VENTILATION_LABELS_SHORT[ventilation_level].lower()
            support_label = support_label.removeprefix("on ")
            clauses.append(f"respiratory support was escalated to {support_label}")

        if ventilation_level is not None:
            previous_ventilation = ventilation_level

        # New-onset abnormal vitals (reported once, the day they
        # first appear, not restated on every day they persist).
        current_abnormal = {
            label for flag, label in _ABNORMAL_VITAL_FLAGS if nss and getattr(nss, flag, None)
        }
        new_abnormal = current_abnormal - previous_abnormal

        if new_abnormal:
            verb = "was" if len(new_abnormal) == 1 else "were"
            clauses.append(f"new {_list_join(sorted(new_abnormal))} {verb} noted")

        previous_abnormal = current_abnormal

        # General condition change.
        current_general_condition = _field(nss, "general_condition") if nss else None

        if (
            current_general_condition
            and previous_general_condition
            and current_general_condition.lower() != previous_general_condition.lower()
        ):
            clauses.append(
                f"general condition changed from {previous_general_condition.lower()} "
                f"to {current_general_condition.lower()}"
            )

        if current_general_condition:
            previous_general_condition = current_general_condition

        # Sepsis impression progression.
        sepsis_text = _field(nss, "sepsis_impression") if nss else None
        sepsis_rank = _sepsis_rank(sepsis_text) if sepsis_text else None
        sepsis_clause = None

        if (
            sepsis_rank is not None
            and previous_sepsis_rank is not None
            and sepsis_rank > previous_sepsis_rank
        ):
            sepsis_clause = f"sepsis impression progressed to {_SEPSIS_LABELS[sepsis_rank]}"
            clauses.append(sepsis_clause)

        if sepsis_rank is not None:
            previous_sepsis_rank = sepsis_rank

        # Treatment (medication regimen) change — only once a prior
        # regimen exists to compare against, so the first-ever
        # recorded medication list reads as the starting regimen, not
        # as an addition to it.
        current_medications = _medication_items(_field(dcm, "medications")) if dcm else {}
        new_medication_keys = set(current_medications) - set(previous_medications)

        if new_medication_keys and previous_medications:
            new_items = [current_medications[key] for key in sorted(new_medication_keys)]
            verb = "was" if len(new_items) == 1 else "were"
            clauses.append(f"{_list_join(new_items)} {verb} added to the medication regimen")

        if current_medications:
            previous_medications = current_medications

        # Weight trend.
        if dcm and dcm.weight is not None and previous_weight is not None:

            delta = dcm.weight - previous_weight

            # Only a clinically noticeable change is worth narrating
            # (day-to-day measurement noise otherwise adds no signal);
            # the threshold does not alter or discard the documented
            # weight itself.
            if abs(delta) >= 50:
                direction = "decreased" if delta < 0 else "increased"
                clauses.append(f"weight {direction} to {_clean(dcm.weight)} g")

        if dcm and dcm.weight is not None:
            previous_weight = dcm.weight

        status = _field(dcm, "vital_status") if dcm else None

        if status and "dead" in status.lower():
            clauses.append("death was recorded")

        if not clauses:
            continue

        sepsis_only = clauses == [sepsis_clause] if sepsis_clause else False

        if sepsis_only and event_sentence_started and not is_last_day:
            # Fold a lone further sepsis-impression step onto the
            # previous per-day event sentence instead of opening a
            # new "On Day N" sentence for it. Varies the connector so
            # a run of several such days doesn't read as
            # "while ..., while ..., while ...".
            connector = "and further" if last_append_was_merge else "while"
            previous_sentence = sentences[-1].rstrip(".")
            sentences[-1] = f"{previous_sentence}, {connector} {sepsis_clause}."
            last_append_was_merge = True
            continue

        if display_date:
            opener = (
                f"By the final documented assessment (Day {day_number}, {display_date})"
                if is_last_day
                else f"On Day {day_number} ({display_date})"
            )
        else:
            opener = (
                f"By the final documented assessment (Day {day_number})"
                if is_last_day
                else f"On Day {day_number}"
            )

        sentences.append(f"{opener}, {_clause_join(clauses)}.")
        event_sentence_started = True
        last_append_was_merge = False

    return " ".join(sentences) if sentences else None


def _section_clinical_course(document: DocxDocument, days: list) -> None:

    _section_heading(document, "3. Clinical Course")

    narrative = _clinical_course_narrative(days)

    _paragraph(
        document,
        narrative or "No Daily Clinical Monitoring or Neonatal Sepsis Screening records were available for this case.",
    )


# --------------------------------------------------------------
# Section 4 — Investigations
# --------------------------------------------------------------

_LAB_PANELS: list[tuple[str, list[tuple[str, str, str]]]] = [
    (
        "CBC",
        [
            ("hemoglobin", "Hemoglobin", "g/dL"),
            ("tlc", "TLC", "/mm³"),
            ("platelet_count", "Platelet count", "/mm³"),
            ("anc", "ANC", "x10³/mm³"),
        ],
    ),
    (
        "LFT",
        [
            ("total_bilirubin", "Total bilirubin", "mg/dL"),
            ("direct_bilirubin", "Direct bilirubin", "mg/dL"),
            ("indirect_bilirubin", "Indirect bilirubin", "mg/dL"),
            ("total_protein", "Total protein", "g/dL"),
            ("albumin", "Albumin", "g/dL"),
            ("globulin", "Globulin", "g/dL"),
            ("alkaline_phosphatase", "Alkaline phosphatase", "IU/L"),
        ],
    ),
    (
        "Electrolytes / KFT",
        [
            ("sodium", "Sodium", "meq/L"),
            ("potassium", "Potassium", "meq/L"),
            ("ionic_calcium", "Ionic calcium", "mg/dL"),
            ("urea", "Urea", "mg/dL"),
            ("creatinine", "Creatinine", "mg/dL"),
        ],
    ),
    (
        "Inflammatory markers",
        [("crp", "CRP", "mg/L")],
    ),
]


def _lab_panel_sentence(panel_name: str, fields: list[tuple[str, str, str]], record: LabData) -> str | None:

    bits = []

    for field_name, label, unit in fields:
        value = _field(record, field_name)
        if value:
            bits.append(f"{label} {value} {unit}")

    return f"{panel_name} — " + ", ".join(bits) if bits else None


_TREND_FIELDS: list[tuple[str, str, str]] = [
    ("tlc", "TLC", "/mm³"),
    ("hemoglobin", "Hb", "g/dL"),
    ("platelet_count", "platelet count", "/mm³"),
    ("anc", "ANC", "x10³/mm³"),
    ("total_bilirubin", "total bilirubin", "mg/dL"),
    ("crp", "CRP", "mg/L"),
    ("blood_glucose", "blood glucose", "mg/dL"),
]


def _section_investigations(document: DocxDocument, summary: CaseSummary) -> None:
    """
    A compact paragraph of documented laboratory findings — standalone
    lab reports plus the day-wise-inline serial values, combined into
    one reading rather than a table, since a typical case has too few
    dated entries for a table to earn its space. Values are shown
    exactly as extracted; a value that looks clinically unusual (e.g.
    an implausible reading, or two records using different units/
    scales for the same test) is still reproduced verbatim, never
    corrected, normalized, or omitted. No trend, percentage, or
    interpretation is calculated — only the documented values and
    their dates are stated.
    """

    _section_heading(document, "4. Investigations")

    sentences: list[str] = []

    # A safe, evidence-based cross-reference (correlating two already
    # -extracted facts) rather than an invented clinical threshold.
    phototherapy_noted = any(
        "photo" in (_field(dcm, "remarks") or "").lower() for dcm in summary.dcm
    )

    for lab in summary.lab:

        panel_sentences = [
            sentence
            for panel_name, fields in _LAB_PANELS
            if (sentence := _lab_panel_sentence(panel_name, fields, lab))
        ]

        if not panel_sentences:
            continue

        header_bits = []

        if _field(lab, "test_date"):
            header_bits.append(_field(lab, "test_date"))

        if _field(lab, "lab_accession_no"):
            header_bits.append(f"Reg/Ref {_field(lab, 'lab_accession_no')}")

        intro = "Laboratory investigations"

        if header_bits:
            intro += f" dated {', '.join(header_bits)}"

        text = intro + " showed: " + "; ".join(panel_sentences) + "."

        if phototherapy_noted and _field(lab, "total_bilirubin"):
            text += " Phototherapy was noted in the clinical record around this time."

        sentences.append(text)

    # Serial values documented inline on the NSS form, one clause per
    # dated record, folded into the same paragraph rather than a
    # separate table.
    serial_bits = []

    for nss in summary.nss:

        date = _field(nss, "assessment_date")

        if not date:
            continue

        value_bits = []

        for field_name, label, unit in _TREND_FIELDS:
            value = _field(nss, field_name)
            if value:
                value_bits.append(f"{label} {value} {unit}")

        if value_bits:
            serial_bits.append(f"on {date}, {_list_join(value_bits)}")

    if serial_bits:
        sentences.append("Serial values documented during monitoring: " + "; ".join(serial_bits) + ".")

    if sentences:
        _paragraph(document, " ".join(sentences))
    else:
        _paragraph(document, "No laboratory investigation results were available for this case.")


# --------------------------------------------------------------
# Section 4 — In-hospital Death & Outcome
# --------------------------------------------------------------


def _last_known_status(nss_records: list[NSSData]) -> str | None:

    dated = [(record, _parse_date(record.assessment_date)) for record in nss_records]
    dated_only = [(record, d) for record, d in dated if d is not None]

    if dated_only:
        dated_only.sort(key=lambda pair: pair[1])
        candidate = dated_only[-1][0]
    elif nss_records:
        candidate = nss_records[-1]
    else:
        return None

    bits = []

    # Lowercased to match the embedded-value casing already used by
    # _progression_narrative / _terminal_course_summary for the same
    # two fields — the source value itself is untouched (still
    # available, original case, via CaseSummary/the day-wise table).
    if _field(candidate, "general_condition"):
        bits.append(f"general condition was documented as {_field(candidate, 'general_condition').lower()}")

    if _field(candidate, "sepsis_impression"):
        bits.append(f"sepsis impression was documented as {_field(candidate, 'sepsis_impression').lower()}")

    if not bits:
        return None

    date_bit = f" as of {candidate.assessment_date}" if candidate.assessment_date else ""

    return f"At the most recent documented Neonatal Sepsis Screening assessment{date_bit}, the infant's {'; '.join(bits)}."


def _terminal_course_summary(days: list) -> str | None:
    """
    A focused, single-day snapshot of the admission's final
    documented status — deliberately narrower than the Day-wise
    Clinical Course (the authoritative daily record of every day) and
    the Progression Narrative (which already synthesizes how the
    trajectory reached this point over multiple days). Restating more
    than the final day here would just repeat, in a third phrasing,
    facts those two sections already cover. This section's only job
    is: where did things stand on the last documented day. Built only
    from fields already present on that record; never a cause.

    Neonatal Sepsis Screening is not always documented on the very
    last day of an admission (routine screening can stop while Daily
    Clinical Monitoring continues) — if the final day has no NSS
    record of its own, this falls back to the most recently
    documented respiratory-support/sepsis status from an earlier day,
    explicitly labeled as such with its own date, rather than this
    snapshot going silent on exactly the day it matters most.
    """

    if not days:
        return None

    day_number, display_date, dcm, nss = days[-1]

    last_ventilation_level: int | None = None
    last_ventilation_ref: tuple[int, str | None] | None = None
    last_sepsis_text: str | None = None
    last_sepsis_ref: tuple[int, str | None] | None = None

    for prior_day_number, prior_display_date, _prior_dcm, prior_nss in days[:-1]:

        ventilation_level = _ventilation_level(prior_nss)

        if ventilation_level is not None:
            last_ventilation_level, last_ventilation_ref = ventilation_level, (prior_day_number, prior_display_date)

        sepsis_text = _field(prior_nss, "sepsis_impression") if prior_nss else None

        if sepsis_text:
            last_sepsis_text, last_sepsis_ref = sepsis_text, (prior_day_number, prior_display_date)

    bits = []

    if dcm and _field(dcm, "vital_status"):
        bits.append(_field(dcm, "vital_status").lower())

    ventilation_level = _ventilation_level(nss)

    if ventilation_level is not None:
        bits.append(_VENTILATION_LABELS_SHORT[ventilation_level].lower())
    elif last_ventilation_level is not None:
        ref_day, ref_date = last_ventilation_ref
        ref_bit = f" (Day {ref_day}{f', {ref_date}' if ref_date else ''})"
        support_label = _VENTILATION_LABELS_SHORT[last_ventilation_level].lower()
        bits.append(f"last documented respiratory support{ref_bit}: {support_label}")

    sepsis_text = _field(nss, "sepsis_impression") if nss else None

    if sepsis_text:
        bits.append(f"sepsis impression {sepsis_text.lower()}")
    elif last_sepsis_text:
        ref_day, ref_date = last_sepsis_ref
        ref_bit = f" (Day {ref_day}{f', {ref_date}' if ref_date else ''})"
        bits.append(f"last documented sepsis impression{ref_bit}: {last_sepsis_text.lower()}")

    if not bits:
        return None

    date_bit = f", {display_date}" if display_date else ""

    return (
        f"On the final documented day of this admission (Day {day_number}{date_bit}), "
        "the infant was documented as " + _list_join(bits) + "."
    )


def _section_outcome(
    document: DocxDocument,
    summary: CaseSummary,
    clinical_report,
    days: list,
) -> None:

    _section_heading(document, "5. Outcome")

    sentences: list[str] = []

    terminal_course = _terminal_course_summary(days)

    if terminal_course:
        sentences.append(terminal_course)

    narrative = _clean(getattr(clinical_report, "final_outcome", None)) if clinical_report else None

    if narrative:
        sentences.append(narrative)

    sae = summary.sae

    if sae:

        onset_bits = []

        if _field(sae, "event_date"):
            onset_bits.append(f"with an onset date of {_field(sae, 'event_date')}")

        diagnosis = _clean_remark(_field(sae, "diagnosis"))

        if diagnosis:
            onset_bits.append(f"diagnosed as {diagnosis}")

        if onset_bits:
            sentences.append("A Serious Adverse Event was recorded " + " and ".join(onset_bits) + ".")

        event_description = _clean_remark(_field(sae, "event_description"))

        if event_description:
            sentences.append(event_description)

        if _field(sae, "seriousness"):
            sentences.append(f"The event was classified as {_field(sae, 'seriousness')}.")

        outcome_bits = []

        cause_of_death = _clean_remark(_field(sae, "cause_of_death"))

        if cause_of_death:
            outcome_bits.append(f"the cause of death was documented as {cause_of_death}")

        if _field(sae, "outcome"):
            outcome_bits.append(f"the documented outcome was {_field(sae, 'outcome')}")

        if outcome_bits:
            combined = " and ".join(outcome_bits)
            sentences.append(combined[0].upper() + combined[1:] + ".")

        # Admission/discharge dates are admission logistics, not
        # outcome facts — their one primary home is the Case Summary
        # section. Reporter name/notification date are regulatory
        # metadata (who filed the report, when), not clinical
        # information, and are intentionally not reproduced in a
        # concise clinical case narrative.

    else:
        sentences.append(
            "No dedicated Serious Adverse Event (SAE) form was identified for this "
            "case in the source document."
        )

    # Only fall back to the last routine monitoring status if the
    # case doesn't already have a clearer outcome statement above —
    # otherwise a pre-event NSS assessment reads as a confusing
    # non-sequitur after an outcome (e.g. death) has just been stated.
    #
    # This also covers terminal_course itself: both it and
    # _last_known_status ultimately read the same NSS records for
    # "the most recent documented sepsis impression" — when
    # terminal_course already states one, appending this fallback
    # would only restate that same fact under a different sentence
    # template. It would NOT be pure redundancy if terminal_course
    # happened to say nothing about sepsis impression (e.g. only a
    # final vital_status with no NSS coverage near the end of the
    # admission at all) — in that case this fallback still adds
    # information terminal_course didn't, so it still fires. A
    # general_condition fact this fallback would otherwise add is not
    # lost when suppressed: it already has its authoritative home in
    # the Day-wise table and, on any day it changes, the Progression
    # Narrative.
    terminal_course_states_sepsis = bool(terminal_course) and "sepsis impression" in terminal_course.lower()

    if (
        not narrative
        and not (sae and _field(sae, "outcome"))
        and not terminal_course_states_sepsis
    ):
        last_status = _last_known_status(summary.nss)
        if last_status:
            sentences.append(last_status)

    _paragraph(document, _join_sentences(sentences))


def _confidentiality_footer(document: DocxDocument) -> None:

    _paragraph(document, "Confidential — ICMR Emollient Study", italic=True)


# --------------------------------------------------------------
# Builder
# --------------------------------------------------------------


class DocumentBuilder:

    def build(
        self,
        summary: CaseSummary,
        output_file: str | Path,
        clinical_report: "ClinicalReport | None" = None,
    ) -> None:

        document = create_document()

        days = _group_by_day(summary.dcm, summary.nss)

        _title_block(document, summary)
        _identification_block(document, summary)

        _section_case_summary(document, summary)
        _section_maternal_history(document, summary, clinical_report)
        _section_clinical_course(document, days)
        _section_investigations(document, summary)
        _section_outcome(document, summary, clinical_report, days)

        _confidentiality_footer(document)

        document.save(str(output_file))
